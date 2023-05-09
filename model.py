from dataclasses import dataclass, asdict
import openpyxl
import json
import os
import datetime
import fpdf
from docx import Document


v0_1_template_constants = {
    'COMMENT_BANK_HEADER_ROWS': 3,
    'REPORT_DATA_HEADER_ROWS': 3,
    'REPORT_METADATA_HEADER_ROWS': 2,

    'STUDENT_LASTNAME_COL': 1,
    'STUDENT_FIRSTNAME_COL': 2,
    'STUDENT_GROUP_COL': 3,
    'STUDENT_NOTES_COL': 4,

    'RAW_COMMENT_COL': 5,
    'COMPILED_COMMENT_COL': 6,

    'DATA_VALUES_START_COL': 7,
    'DATA_VALUES_LABELS_ROW': 3
    
}


@dataclass
class Student:
    firstname: str
    lastname: str
    group: str
    notes: str


@dataclass
class Report:
    id: int
    raw_comment: str
    compiled_comment: str
    student: Student
    complete: bool
    data_values: dict


@dataclass
class Comment:
    id: int
    label: str
    text: str
    category: str


class LoadExcelTemplateException(Exception):
    pass

class SaveExcelTemplateException(Exception):
    pass

class ExportReportSetException(Exception):
    pass

class ReportSet:

    def __init__(self, reports=[], comment_bank=[]):
        self.reports = reports
        self.comment_bank = comment_bank

    def reports_as_json(self):
        # TODO: Check if this is correct - should it just be {'reports': ... }?
        return json.dumps({self.reports: [asdict(r) for r in self.reports]})

    def comment_bank_as_json(self):
        return json.dumps({'comment_bank': [asdict(c) for c in self.comment_bank]})

    def report_set_as_json(self):
        return json.dumps(
            {
                'reports':[asdict(r) for r in self.reports],
                'comment_bank': [asdict(c) for c in self.comment_bank]
            }
        )

    def save_to_excel_template(self, path_to_saved_file=None):

        if len(self.reports) == 0:
            raise SaveExcelTemplateException("No report data to save!")

        if path_to_saved_file is None:
            path_to_saved_file = os.path.join('temp', f"ReportWriterExport_{datetime.datetime.strftime(datetime.datetime.now(),'%Y-%m-%d_%H%M%S.xlsx')}")

        doc = openpyxl.load_workbook(os.path.join('resources', 'template.xlsx'))
        template_version = doc["Info"]["B4"].value

        if template_version == 0.1:
            template_constants = v0_1_template_constants

        report_data = doc['Data entry']
        report_metadata = doc['Metadata']

        # Add Data Value labels
        for i, label in enumerate(self.reports[0].data_values.keys()):
            report_data.cell(template_constants['DATA_VALUES_LABELS_ROW'],
                             template_constants['DATA_VALUES_START_COL'] + i).value = label

        # Add report data
        for i, report in enumerate(self.reports):
            row = template_constants['REPORT_DATA_HEADER_ROWS'] + 1 + i
            metadata_row = template_constants['REPORT_METADATA_HEADER_ROWS'] + 1 + i
            report_data.cell(row, 1).value = report.student.lastname
            report_data.cell(row, 2).value = report.student.firstname
            report_data.cell(row, 3).value = report.student.group
            report_data.cell(row, 4).value = report.student.notes
            report_data.cell(row, 5).value = report.raw_comment
            report_data.cell(row, 6).value = report.compiled_comment
            #report_data.cell(row, 7).value = report.compiled_comment

            # Add data values data
            for j, label in enumerate(report.data_values.keys()):
                report_data.cell(row, template_constants['DATA_VALUES_START_COL'] + j).value = report.data_values[label]

            # Add metadata
            report_metadata.cell(metadata_row, 1).value = i + 1
            report_metadata.cell(metadata_row, 2).value = report.raw_comment
            report_metadata.cell(metadata_row, 3).value = report.compiled_comment
            report_metadata.cell(metadata_row, 4).value = 'Y' if report.complete else 'N'

        # Add comment bank
        comment_data = doc['Comment bank']

        # TODO: Update to use new comment bank list
        for i, comment in enumerate(self.comment_bank):
            row = template_constants['COMMENT_BANK_HEADER_ROWS'] + 1 + i
            comment_data.cell(row, 1).value = comment.label
            comment_data.cell(row, 2).value = comment.text
            comment_data.cell(row, 3).value = comment.category

        # Save finished document
        doc.save(path_to_saved_file)

        # Return path to saved file
        return path_to_saved_file

    def export_to_txt(self, path_to_export_file=None):

        if len(self.reports) == 0:
            raise SaveExcelTemplateException("No report data to save!")

        if path_to_export_file == None:
            path_to_export_file = os.path.join('temp', f"ReportWriterExport_{datetime.datetime.strftime(datetime.datetime.now(),'%Y-%m-%d_%H%M%S.txt')}")

        with open(path_to_export_file, "w") as f:
            for report in self.reports:
                f.write(f"{report.student.firstname} {report.student.lastname} ({report.student.group})\n")
                f.write(f"{'=' * 80}\n")
                f.write(f"{report.compiled_comment}\n")
                f.write(f"{'-' * 80}\n\n")

        # Return path to saved file
        return path_to_export_file

    def export_to_pdf(self, path_to_export_file=None):
        if len(self.reports) == 0:
            raise SaveExcelTemplateException("No report data to save!")

        if path_to_export_file == None:
            path_to_export_file = os.path.join('temp', f"ReportWriterExport_{datetime.datetime.strftime(datetime.datetime.now(),'%Y-%m-%d_%H%M%S.pdf')}")

        pdf = fpdf.FPDF('P', 'mm', 'A4')
        pdf.add_page()

        for report in self.reports:
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, f"{report.student.firstname} {report.student.lastname} ({report.student.group})")
            pdf.ln(10)
            pdf.set_font('Arial', '', 12)
            pdf.multi_cell(0, 5, f"{report.compiled_comment}")
            pdf.ln(8)

        pdf.output(path_to_export_file, 'F')

        # Return path to saved file
        return path_to_export_file

    def export_to_word(self, path_to_export_file=None):

        if len(self.reports) == 0:
            raise SaveExcelTemplateException("No report data to save!")

        if path_to_export_file == None:
            path_to_export_file = os.path.join('temp', f"ReportWriterExport_{datetime.datetime.strftime(datetime.datetime.now(),'%Y-%m-%d_%H%M%S.docx')}")

        doc = Document()

        for report in self.reports:
            doc.add_heading(f"{report.student.firstname} {report.student.lastname} ({report.student.group})")
            doc.add_paragraph(f"{report.compiled_comment}")

        doc.save(path_to_export_file)

        # Return path to saved file
        return path_to_export_file

    @staticmethod
    def load_from_json(rs_json):
        rs = json.loads(rs_json)

        reports = [
            Report(r['id'],
                   r['raw_comment'],
                   r['compiled_comment'],
                   Student(
                       r['student']['firstname'],
                       r['student']['lastname'],
                       r['student']['group'],
                       r['student']['notes']
                   ),
                   r['complete'],
                   r['data_values'])
            for r in rs['reports']
        ]

        comment_bank = [
            Comment(c['id'], c['label'], c['text'], c['category']) for c in rs['comment_bank']
        ]

        return ReportSet(reports=reports, comment_bank=comment_bank)

    @staticmethod
    def load_from_excel_template(template_path):

        try:
            xl_template = openpyxl.load_workbook(template_path)
            template_version = xl_template["Info"]["B4"].value

            if template_version == 0.1:
                template_constants = v0_1_template_constants

            # Build comment bank
            comment_data = xl_template["Comment bank"]

            # Count the number of rows until a blank is found in the comment bank sheet
            row = 1
            while (comment_data.cell(row, 1).value is not None):
                row += 1
            comment_count = row - template_constants['COMMENT_BANK_HEADER_ROWS'] - 1

            # TODO: Update to use new comment bank list
            comment_bank = []

            for i in range(1, comment_count + 1):
                if comment_data.cell(template_constants['COMMENT_BANK_HEADER_ROWS'] + i, 1).value is not None:
                    label = comment_data.cell(template_constants['COMMENT_BANK_HEADER_ROWS'] + i, 1).value
                    text = comment_data.cell(template_constants['COMMENT_BANK_HEADER_ROWS'] + i, 2).value
                    category = comment_data.cell(template_constants['COMMENT_BANK_HEADER_ROWS'] + i, 3).value

                    comment_bank.append(Comment(i,label, text, category))

            # Build list of reports
            report_data = xl_template["Data entry"]
            report_metadata = xl_template["Metadata"]
            report_count = report_data.max_row - template_constants['REPORT_DATA_HEADER_ROWS']

            if report_count == 0:
                raise LoadExcelTemplateException("Excel sheet contains no student rows! Please add at least one student before continuing.")

            # data_value_count = report_data.max_column - (template_constants['DATA_VALUES_START_COL'] - 1)

            data_value_count = 0
            col = template_constants['DATA_VALUES_START_COL']
            while report_data.cell(template_constants['DATA_VALUES_LABELS_ROW'],col).value != None:
                x = report_data.cell(template_constants['DATA_VALUES_LABELS_ROW'],col).value
                data_value_count += 1
                col += 1

            reports = []

            for row in range(template_constants['REPORT_DATA_HEADER_ROWS'] + 1,
                             template_constants['REPORT_DATA_HEADER_ROWS'] + report_count + 1):

                report_data_values = {}

                for col in range(template_constants['DATA_VALUES_START_COL'],
                                 template_constants['DATA_VALUES_START_COL'] + data_value_count):
                    report_data_values[report_data.cell(template_constants['DATA_VALUES_LABELS_ROW'], col).value] = \
                        report_data.cell(row, col).value

                # Compare report_data's compiled report to metadata - if they differ, take the report_data compiled
                # report as both raw and compiled report comments, as the user has overwritten.

                raw_comment = report_data.cell(row, template_constants['RAW_COMMENT_COL']).value
                if report_data.cell(row, template_constants['COMPILED_COMMENT_COL']).value != \
                        report_metadata.cell(row - 1, 3).value:
                    raw_comment = report_data.cell(row, template_constants['COMPILED_COMMENT_COL']).value

                reports.append(
                    Report(
                        id=row - template_constants['REPORT_DATA_HEADER_ROWS'],
                        # raw_comment=report_data.cell(row, template_constants['RAW_COMMENT_COL']).value,
                        raw_comment=raw_comment,
                        compiled_comment=report_data.cell(row, template_constants['COMPILED_COMMENT_COL']).value,
                        student=Student(
                            firstname=report_data.cell(row, template_constants['STUDENT_FIRSTNAME_COL']).value,
                            lastname=report_data.cell(row, template_constants['STUDENT_LASTNAME_COL']).value,
                            group=report_data.cell(row, template_constants['STUDENT_GROUP_COL']).value,
                            notes=report_data.cell(row, template_constants['STUDENT_NOTES_COL']).value
                        ),
                        complete=report_metadata.cell(row - 1, 4).value == 'Y', # Metadata sheet has one fewer heading rows than report sheet, hence row - 1
                        data_values=report_data_values
                    )
                )

            return ReportSet(reports, comment_bank)

        except openpyxl.utils.exceptions.InvalidFileException as e:
            raise LoadExcelTemplateException(str(e))

        except FileNotFoundError as e:
            raise LoadExcelTemplateException(str(e))

        except Exception as e:
            raise e


class AEASReportSetConverter:

    # Compatible with January 2022 Beta B AEAS version

    ASSESSMENT_PERSPECTIVE_LABLES_START_ROW = 11
    ASSESSMENT_PERSPECTIVE_1_COL = 4
    ASSESSMENT_PERSPECTIVE_2_COL = 5

    STUDENT_DATA_FIRST_ROW = 5
    STUDENT_QUESTION_MAX_MARKS_ROW = 4
    STUDENT_NAME_COL = 1
    STUDENT_GROUP_COL = 2
    STUDENT_NOTES_COL = 3
    STUDENT_QUESTION_DATA_START_COL = 4


    class LoadAEASError(Exception):
        pass

    @staticmethod
    def getReportSetFromAEAS(path_to_aeas):

        try:
            rs = ReportSet()
            aeas = openpyxl.load_workbook(path_to_aeas)
            setup_page = aeas['Setup']
            DES = aeas['Data Entry Sheet']

            # Build data values by iterating through Setup page's assessment perspective labels

            data_value_labels = {}

            # Get Assessment Perspective 1 labels
            row = AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW
            ass_perspective_labels_count = 0

            while setup_page.cell(row, AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_1_COL).value != None:
                ass_perspective_labels_count += 1
                row += 1

            for row in range(AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW, AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW + ass_perspective_labels_count):
                ass_perspective = setup_page.cell(row, AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_1_COL).value
                if ass_perspective not in data_value_labels:
                    data_value_labels[ass_perspective] = []
                q_ref = (row - AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW) + 1
                data_value_labels[ass_perspective].append(q_ref)

            # Get Assessment Perspective 2 labels
            row = AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW
            ass_perspective_labels_count = 0

            while setup_page.cell(row, AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_2_COL).value != None:
                ass_perspective_labels_count += 1
                row += 1

            for row in range(AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW, AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW + ass_perspective_labels_count):
                ass_perspective = setup_page.cell(row, AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_2_COL).value
                if ass_perspective not in data_value_labels:
                    data_value_labels[ass_perspective] = []
                q_ref = (row - AEASReportSetConverter.ASSESSMENT_PERSPECTIVE_LABLES_START_ROW) + 1
                data_value_labels[ass_perspective].append(q_ref)


            # Get number of students from data entry sheet
            row = AEASReportSetConverter.STUDENT_DATA_FIRST_ROW
            student_count = 0
            while (DES.cell(row, 1).value != None):
                student_count += 1
                row += 1

            # For each student, create add a report and iterate through their question data to extract
            # % for each data value (assessment perspective)


            for i in range(student_count):
                row = AEASReportSetConverter.STUDENT_DATA_FIRST_ROW + i
                name = DES.cell(row, AEASReportSetConverter.STUDENT_NAME_COL).value

                fn = name.split(" ")[0]
                ln = ''.join(name.split(" ")[1:])
                group = DES.cell(row, AEASReportSetConverter.STUDENT_GROUP_COL).value
                notes = DES.cell(row, AEASReportSetConverter.STUDENT_NOTES_COL).value
                # group = DES.cell(row, AEASReportSetConverter.STUDENT_GROUP_COL).value if DES.cell(row, AEASReportSetConverter.STUDENT_GROUP_COL).value is not None else ''
                # notes = DES.cell(row, AEASReportSetConverter.STUDENT_NOTES_COL).value if DES.cell(row, AEASReportSetConverter.STUDENT_NOTES_COL).value is not None else ''
                dvs = {}

                # Find the relevant data for that student's data values
                for key, qrefs in data_value_labels.items():
                    total = 0
                    max = 0

                    for q in qrefs:
                        try:
                            total += float(DES.cell(row, AEASReportSetConverter.STUDENT_QUESTION_DATA_START_COL - 1 + q).value)
                            max += float(DES.cell(AEASReportSetConverter.STUDENT_QUESTION_MAX_MARKS_ROW, AEASReportSetConverter.STUDENT_QUESTION_DATA_START_COL - 1 + q).value)
                        except ValueError:
                            pass
                    try:
                        dvs[key] = f"{(total * 100 / max):0.1f}"
                    except ZeroDivisionError:
                        dvs[key] = 0

                report = Report(i + 1, "", "", Student(fn, ln, group, notes), False, dvs)

                rs.reports.append(report)

            # Export ReportWriter Template from report set
            # rs.save_to_excel_template()
            return rs

        except openpyxl.utils.exceptions.InvalidFileException as e:
            raise AEASReportSetConverter.LoadAEASError(str(e))

        except FileNotFoundError as e:
            raise AEASReportSetConverter.LoadAEASError(str(e))

        except Exception as e:
            raise e


if __name__ == '__main__':
    pass

